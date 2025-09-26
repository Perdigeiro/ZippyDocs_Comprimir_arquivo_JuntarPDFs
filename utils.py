"""
Variáveis globais introduzidas para resolver problemas de serialização no multiprocessing,
como o erro "cannot pickle '_tkinter.tkapp' object". Ao mover a lógica de conversão para funções
independentes, evita a dependência de objetos Tkinter associados á instância da interface,
permitindo que o código funcione corretamente em subprocessos com segurança mantida.
"""

import pdfplumber                               # Para extração de tabelas de PDFs
import pandas as pd                             # Para manipulação de DataFrames
import os                                       # Para manipulação de caminhos e arquivos
from pdf2docx import Converter                  # Para conversão de PDF para Word
from pypdf import PdfReader, PdfWriter          # Para manipulação de PDFs
from pdf2image import convert_from_path         # Para conversão de PDF para imagens
import zipfile                                  # Para criação/extração de arquivos ZIP
from docx import Document                       # Para manipulação de arquivos Word
import re                                       # Para expressões regulares
from docx2pdf import convert as docx2pdf_convert   # Para conversão de DOCX para PDF
import mimetypes                                # Para manipulação de tipos MIME
import logging                                # Para registro de logs
import tempfile                                 # Para criação de arquivos temporários
import traceback                             # Para captura de rastreamentos de erros
import chardet                                  # Para detecção de codificação de arquivos
import sys                                      # Para manipulação de sistema
import uuid                                     # Para geração de identificadores únicos
import magic                                    # Para detecção de tipos MIME com python-magic
import re                                      # Para expressões regulares
import hashlib                                 # Para cálculo de hash de arquivos
import subprocess
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4


# Criar pasta C:\Temp\Logs_conversor se não existir
if not os.path.exists(r'C:\Temp\Logs_conversor'):
    os.makedirs(r'C:\Temp\Logs_conversor')

# Configurar logger para criar arquivo por erro
logger = logging.getLogger("teste")
logger.setLevel(logging.DEBUG)

if not logger.handlers:
    # Usar nome fixo erro_conversor.log
    nome_arquivo = "erro_conversor.log"
    caminho_arquivo = os.path.join(r'C:\Temp\Logs_conversor', nome_arquivo)
    handler = logging.FileHandler(caminho_arquivo, encoding='utf-8')
    formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)


def obter_caminho_poppler(): # Função para obter o caminho do Poppler
    """Retorna o caminho do Poppler, ajustando para o ambiente do PyInstaller."""
    if hasattr(sys, '_MEIPASS'):  # Verifica se está rodando como executável
        return os.path.join(sys._MEIPASS, 'poppler')  # Caminho dentro do executável
    return r'C:\poppler-25.07.0\Library\bin'  # Caminho local para desenvolvimento

def obter_caminho_grhostscript(): # Função para obter o caminho do Poppler
    """Retorna o caminho do Grhostscript, ajustando para o ambiente do PyInstaller."""
    if hasattr(sys, '_MEIPASS'):  # Verifica se está rodando como executável
        return os.path.join(sys._MEIPASS, 'gs')  # Caminho dentro do executável
    return r'C:\Program Files\gs\gs10.06.0\bin' # Caminho local para desenvolvimento



def wrapper(q, func, *args, **kwargs):
    try:
        q.put((True, func(*args, **kwargs)))
    except Exception as e:
        q.put((False, e))

def juntar_pdfs_worker(lista_arquivos, arquivo_saida):
    """Junta PDFs convertendo cada página para imagem (garantia máxima de visualização)."""
    c = canvas.Canvas(arquivo_saida, pagesize=A4)
    for arquivo in lista_arquivos:
        imagens = convert_from_path(arquivo, poppler_path=obter_caminho_poppler())
        for img in imagens:
            img_width, img_height = img.size
            a4_width, a4_height = A4
            # Redimensiona para caber na página A4
            ratio = min(a4_width / img_width, a4_height / img_height)
            new_width = img_width * ratio
            new_height = img_height * ratio
            img_temp = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.jpg")
            img = img.convert("RGB")
            img.save(img_temp, "JPEG", quality=95)
            c.drawImage(img_temp, (a4_width - new_width) / 2, (a4_height - new_height) / 2, new_width, new_height)
            c.showPage()
            os.remove(img_temp)
    c.save()
    os.chmod(arquivo_saida, 0o600)


def validar_mime(filepath, tipos_aceitos):
    """Valida o tipo MIME real do arquivo usando python-magic, com fallback para mimetypes."""
    try:
        mime = magic.from_file(filepath, mime=True)
        print(f"[DEBUG] MIME detectado: {mime}")
        if mime in tipos_aceitos:
            return True
        # Fallback: tenta pelo mimetypes se magic falhar
        mime_guess, _ = mimetypes.guess_type(filepath)
        print(f"[DEBUG] MIME (fallback mimetypes): {mime_guess}")
        if mime_guess and mime_guess in tipos_aceitos:
            return True
        return False
    except Exception as e:
        print(f"[DEBUG] Erro na detecção MIME: {e}")
        # Fallback direto para mimetypes se magic falhar
        mime_guess, _ = mimetypes.guess_type(filepath)
        print(f"[DEBUG] MIME (fallback mimetypes): {mime_guess}")
        if mime_guess and mime_guess in tipos_aceitos:
            return True
        return False
    
def calcular_sha256(caminho):
    with open(caminho, 'rb') as f:
        return hashlib.sha256(f.read()).hexdigest()
    
def sanitizar_celula_excel(valor):
    """Prefixa com apóstrofo se for potencial fórmula perigosa."""
    if isinstance(valor, str) and valor and valor[0] in ('=', '+', '-', '@'):
        return "'" + valor
    return valor
    

def pdf_para_docx_global(origem, destino):
    """Converte PDF para DOCX de forma independente."""
    try:
        cv = Converter(origem)
        cv.convert(destino)
        cv.close()
        os.chmod(destino, 0o600)
    except Exception as e:
        logger.error(f"Erro ao converter PDF para DOCX: {str(e)}")
        logger.error(traceback.format_exc())


def pdf_para_txt_global(origem, destino):
    """Converte PDF para TXT de forma independente."""
    try:
        with open(origem, 'rb') as f:
            reader = PdfReader(f)
            texto_extraido = ""
            for pagina in reader.pages:
                texto_extraido += pagina.extract_text()
        with open(destino, 'w', encoding='utf-8') as f_out:
            f_out.write(texto_extraido)
        os.chmod(destino, 0o600)
    except Exception as e:
        logger.error(f"Erro ao converter PDF para TXT: {str(e)}")
        logger.error(traceback.format_exc())

def pdf_para_png_global(origem, destino_zip):
    """Converte PDF para PNG e compacta em um ZIP de forma independente."""
    try:
        poppler_path = obter_caminho_poppler()
        imagens = convert_from_path(origem, poppler_path=poppler_path)
        temp_dir = os.path.dirname(destino_zip)
        caminhos_png = []

        for i, img in enumerate(imagens):
            img_path = os.path.join(temp_dir, f"pagina_{i+1}.png")
            img.save(img_path, 'PNG')
            caminhos_png.append(img_path)

        # Compacta os PNGs em um ZIP
        with zipfile.ZipFile(destino_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for img_path in caminhos_png:
                zipf.write(img_path, os.path.basename(img_path))

        # Remove os PNGs temporários
        for img_path in caminhos_png:
            os.remove(img_path)
    except Exception as e:
        logger.error(f"Erro ao converter PDF para PNG: {str(e)}")
        logger.error(traceback.format_exc())

def docx_para_pdf_global(arquivo_origem, arquivo_destino):
    """Converte DOCX para PDF de forma independente e segura."""
    try:
        # --- Segurança: Sanitização do texto do DOCX ---
        doc = Document(arquivo_origem)
        texto_completo = ""
        for para in doc.paragraphs:
            texto = para.text
            # Remove tags HTML/JS
            texto = re.sub(r'<[^>]+>', '', texto)
            # Remove URLs (links suspeitos)
            texto = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', texto)
            # Remove fórmulas perigosas
            if texto and texto[0] in ('=', '+', '-', '@'):
                texto = "'" + texto
            texto_completo += texto + "\n"

        # --- Conversão usando docx2pdf ---
        docx2pdf_convert(arquivo_origem, arquivo_destino)

        # --- Validação MIME do PDF gerado ---
        if not validar_mime(arquivo_destino, ['application/pdf', 'application/x-pdf']):
            os.remove(arquivo_destino)
            raise ValueError("O arquivo gerado não é um PDF válido e foi removido por segurança.")

        # --- Permissões restritas ---
        os.chmod(arquivo_destino, 0o600)

    except Exception as e:
        logger.error(f"Erro ao converter DOCS para PDF: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"Erro seguro na conversão DOCX para PDF: {str(e)}")

def docx_para_txt_global(origem, destino):
    """Converte DOCX para TXT de forma independente."""
    try:
        doc = Document(origem)
        with open(destino, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
        os.chmod(destino, 0o600)
    except Exception as e:
        logger.error(f"Erro ao converter DOCX para TXT: {str(e)}")
        logger.error(traceback.format_exc())

def docx_para_png_global(origem, destino_zip):
    """Converte DOCX para PNG e compacta em um ZIP de forma independente."""
    try:
        # Gera PDF temporário a partir do DOCX
        pdf_temp = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.pdf")
        docx2pdf_convert(origem, pdf_temp)

        # Converte o PDF temporário para imagens PNG
        poppler_path = obter_caminho_poppler()
        imagens = convert_from_path(pdf_temp, poppler_path=poppler_path)
        temp_dir = os.path.dirname(destino_zip)
        caminhos_png = []

        for i, img in enumerate(imagens):
            img_path = os.path.join(temp_dir, f"pagina_{i+1}.png")
            img.save(img_path, "PNG")
            os.chmod(img_path, 0o600)
            caminhos_png.append(img_path)

        # Compacta os PNGs em um ZIP
        with zipfile.ZipFile(destino_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for img_path in caminhos_png:
                zipf.write(img_path, os.path.basename(img_path))

        # Remove os PNGs temporários e o PDF temporário
        for img_path in caminhos_png:
            os.remove(img_path)
        os.remove(pdf_temp)
    except Exception as e:
        logger.error(f"Erro ao converter DOCX para PNG: {str(e)}")
        logger.error(traceback.format_exc())


def csv_para_xlsx_global(origem, destino):
    '''Converte CSV para Excel (XLSX) de forma independente, segura e auditável.'''
    try:
        # Detecta a codificação do arquivo CSV
        with open(origem, 'rb') as f:
            resultado = chardet.detect(f.read())
            encoding_detectada = resultado['encoding']

        # Lê o arquivo CSV com a codificação detectada
        df = pd.read_csv(origem, sep=None, engine='python', encoding=encoding_detectada)

        df = df.applymap(sanitizar_celula_excel)

        # Salva o DataFrame como um arquivo Excel
        df.to_excel(destino, index=False, engine='openpyxl')
        os.chmod(destino, 0o600)  # Permissões restritas

    except Exception as e:
        logger.error(f"Erro ao converter CSV para XLSX: {str(e)}")
        logger.error(traceback.format_exc())

def xlsx_para_csv_global(origem, destino):
    try:
        """Converte Excel (XLSX) para CSV de forma independente e segura."""
        df = pd.read_excel(origem)
        df.to_csv(destino, index=False, encoding='utf-8')
        os.chmod(destino, 0o600)
    except Exception as e:
        logger.error(f"Erro ao converter XLSX para CSV: {str(e)}")
        logger.error(traceback.format_exc())


def pdf_para_xlsx_global(origem, destino):
    """
    Converte tabelas de um PDF em um arquivo Excel (.xlsx), com validações de segurança.
    """
    try:
        if not os.path.exists(origem):
            raise FileNotFoundError(f"Arquivo de origem não encontrado: {origem}")

        if not validar_mime(origem, ["application/pdf"]):
            raise ValueError("Tipo de arquivo inválido. Apenas PDFs são permitidos.")

        dados_finais = []
        cabecalhos = None

        with pdfplumber.open(origem) as pdf:
            if not pdf.pages:
                raise ValueError(f"O PDF {origem} não contém páginas.")

            for i, pagina in enumerate(pdf.pages):
                try:
                    tabelas = pagina.extract_tables({
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                        "intersection_y_tolerance": 5,
                        "intersection_x_tolerance": 5,
                        "snap_tolerance": 2,
                        "explicit_vertical_lines": [],
                        "explicit_horizontal_lines": []
                    })

                    for tabela in tabelas or [[]]:
                        for linha in tabela or []:
                            if not linha:
                                continue
                            linha_limpa = [str(c).strip() if c else '' for c in linha]
                            linha_limpa = [c for c in linha_limpa if c]

                            if not linha_limpa:
                                continue

                            # Ignorar rodapés ou trechos irrelevantes
                            linha_texto = ' '.join(linha_limpa).lower()
                            palavras_ignoradas = [
                                'classificado como interno',
                                'usuário',
                                'documento',
                                'grupo fleury',
                                '#',
                                '##'
                            ]
                            if any(p in linha_texto for p in palavras_ignoradas):
                                continue

                            # Detectar cabeçalho
                            if not cabecalhos and not any(re.match(r'\d{2}/\d{2}/\d{4}', c) for c in linha_limpa):
                                cabecalhos = linha_limpa
                                if len(cabecalhos) < 2:
                                    cabecalhos = [f"Coluna_{i}" for i in range(len(linha_limpa))]
                                continue

                            # Ajustar colunas
                            if cabecalhos:
                                num_colunas = len(cabecalhos)
                                if len(linha_limpa) > num_colunas:
                                    linha_limpa = linha_limpa[:num_colunas - 1] + [' '.join(linha_limpa[num_colunas - 1:])]
                                elif len(linha_limpa) < num_colunas:
                                    linha_limpa += [''] * (num_colunas - len(linha_limpa))
                            dados_finais.append(linha_limpa)
                except Exception as page_error:
                    logger.warning(f"Erro na página {i+1}: {str(page_error)}")
                    continue

        if not dados_finais:
            raise ValueError("Nenhum dado válido foi extraído do PDF.")

        colunas = cabecalhos if cabecalhos and len(cabecalhos) == len(dados_finais[0]) else [f"Coluna_{i}" for i in range(len(dados_finais[0]))]

        if any(len(row) != len(colunas) for row in dados_finais):
            raise ValueError("Inconsistência detectada no número de colunas dos dados extraídos.")

        df = pd.DataFrame(dados_finais, columns=colunas)
        df = df.applymap(sanitizar_celula_excel)

        df.to_excel(destino, index=False, engine='openpyxl')
        os.chmod(destino, 0o600)

        # Gravar hash de integridade
        hash_pdf = calcular_sha256(origem)
        with open(destino + ".hash.txt", "w", encoding='utf-8') as f:
            f.write(f"SHA256 do PDF original: {hash_pdf}\n")

    except Exception as e:
        logger.error(f"Erro ao converter PDF para XLSX: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception("Falha ao converter PDF para XLSX.")


def comprimir_pdf(input_pdf, output_pdf):
    """Comprime um PDF usando Ghostscript com caminho dinâmico, sem abrir terminal."""
    ghostscript_dir = obter_caminho_grhostscript()
    ghostscript_exe = os.path.join(ghostscript_dir, "gswin64c.exe")

    if not os.path.exists(ghostscript_exe):
        raise FileNotFoundError(f"Ghostscript não encontrado em: {ghostscript_exe}")

    comando = [
        ghostscript_exe,
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        "-dPDFSETTINGS=/ebook",
        "-dNOPAUSE",
        "-dQUIET",
        "-dBATCH",
        f"-sOutputFile={output_pdf}",
        input_pdf
    ]

    try:
        subprocess.run(
            comando,
            check=True,
            creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
        )
    except Exception as e:
        logger.error(f"Erro ao comprimir PDF: {e}")
        logger.error(traceback.format_exc())

    return output_pdf


if __name__ == "__main__":
    pass
